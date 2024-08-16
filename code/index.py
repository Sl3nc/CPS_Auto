from tkinter import *
from tkinter import ttk
from tkinter import messagebox
from tkinter.filedialog import asksaveasfilename
from num2words import num2words
from docx import Document
from abc import abstractmethod
import keyboard
import re
import os


def enter_press(event):
    if event.keysym == 'Return':
        keyboard.send('tab')
    elif event.keysym == 'Down' or event.keysym == 'Up':
        keyboard.send('space')
        
window = Tk()
window.bind('<KeyRelease>', enter_press)


class Formater: #TODO Formaters
    def cpf_formater(text, var, index, mode):
        #Só recebe valor que passa pelo validador
        valor = text.get()
        if len(valor) == 11:
           valor = valor[:3] + "." + valor[3:6] + "." + valor[6:9] + "-" + valor[9:]
        else:
            valor = valor.replace('.','').replace('-','')
        text.set(valor)

    def cnpj_formater(text, var, index, mode): 
        #Só recebe valor que passa pelo validador
        valor = text.get()
        if len(valor) == 14:
           valor = valor[:2] + "." + valor[2:5] + "." + valor[5:8] + "/" + valor[8:12] + "-" + valor[12:]
        else:
            valor = valor.replace('.','').replace('-','').replace('/','')
        text.set(valor)
    

    def cep_formater(text, var, index, mode): 
        #Só recebe valor que passa pelo validador
        valor = text.get()
        if len(valor) == 8:
           valor = valor[:5] + "-" + valor[5:]
        else:
            valor = valor.replace('-','')
        text.set(valor)

    def rg_formater(text, var, index, mode): 
        #Só recebe valor que passa pelo validador
        valor = text.get()
        if len(valor) == 10:
           valor = valor[:2] + "-" + valor[2:4] + "." + valor[4:7] + "." + valor[7:]
        else:
            valor = valor.replace('.','').replace('-','')
        text.set(valor)
    
    def date_formater(text, var, index, mode): 
        #Só recebe valor que passa pelo validador
        valor = text.get()
        if len(valor) == 8:
           valor = valor[:2] + "/" + valor[2:4] + "/" + valor[4:]
        else:
            valor = valor.replace('/','')
        text.set(valor)

    def valor_formater(text, var, index, mode): 
        #Só recebe valor que passa pelo validador
        valor = text.get()
        if valor[0:2] != 'R$':
            valor = valor.replace('R', '')
            valor = valor.replace('$', '')
            valor = 'R$' + valor
        if ',' not in valor:
            valor = valor + ',00'
        
        text.set(valor)      

class Validator:    #TODO Validators
    def str_validator(text):
        return not text.isdecimal()
    
    def num_validator(text):
        return text.isdecimal()
    
    def cpf_validator(text):
        padrao = r"^[-\d.,/]+$"  # Permite dígitos, ponto, vírgula, hífen e barra
        if len(text) < 15:
            if len(text) >= 12:
                return re.match(padrao, text) is not None
            else:
                return text.isdecimal()
        return False
        
    def cnpj_validator(text):
        padrao = r"^[-\d.,/]+$"  # Permite dígitos, ponto, vírgula, hífen e barra
        if len(text) < 19:
            if len(text) >= 15:
                return re.match(padrao, text) is not None
            else:
                return text.isdecimal()
        return False
    
    def cep_validator(text):
        padrao = r"^[-\d.,/]+$"  # Permite dígitos, ponto, vírgula, hífen e barra
        if len(text) < 10:
            if len(text) >= 9:
                return re.match(padrao, text) is not None
            else:
                return text.isdecimal()
        return False
    
    def rg_validator(text):
        if len(text) < 14:
            return True
        return False
    
    def date_validator(text):
        padrao = r"^[-\d.,/]+$"  # Permite dígitos, ponto, vírgula, hífen e barra
        if len(text) < 11:
            if len(text) >= 9:
                return re.match(padrao, text) is not None
            else:
                return text.isdecimal()
        return False

class File:
    def __init__(self, caminho):
        self.arquivo = Document(f'./code/CPS\'s/CPS {caminho.uppcase()} .docx')    

    def alterar(self, referencias):
        self.mudar_pre_envio()
        
        if self.input_vazio():
            messagebox.showwarning(title='Aviso', message='Existem entradas vazias, favor preencher todas')
            return None
        

        for par in self.arquivo.paragraphs:
            for itens in referencias:
                if par.text.find(itens) != -1:
                    par.text = par.text.replace(itens, referencias[itens].get())


        file = asksaveasfilename(title='Defina o nome e o local onde o arquivo será salvo', filetypes=((".docx","*.docx"),))

        self.doc.save(file+'.docx')

        messagebox.showinfo(title='Aviso', message='Abrindo o arquivo gerado!')

        os.startfile(file+'.docx')

        frame_ativo.destroy()
        self.pageMenu()

class Pages:
    def __init__(self, titulo):
        self.titulo = titulo
        self.file = File(caminho)
        
        @abstractmethod
        def index(self):
            pass
        
        def executar(self):
            self.file.alterar(self.referencias)

class Enterprise(Pages):
    def __init__(self, titulo):
        super().__init__(titulo)
        self.referencias = {
            '$nomeEmp' : StringVar(), #strValidator
            '$ruaEmp' : StringVar(), #strValidator
            '$numEmp' : StringVar(), #intValidator
            '$bairroEmp' : StringVar(),  #strValidator
            '$cepEmp' : StringVar(),  #intValidator
            '$rgEmp' : StringVar(),  #rgValidator
            '$sspEmp' : StringVar(),  #sspValidator
            '$cnpjEmp' : StringVar(),  #cpfValidator
            '$nomeContra' : StringVar(), #strValidator
            '$ruaContra' : StringVar(), #strValidator
            '$numContra' : StringVar(), #intValidator
            '$bairroContra' : StringVar(),  #strValidator
            '$cepContra' : StringVar(),  #intValidator
            '$rgContra' : StringVar(),  #rgValidator
            '$emissorContra' : StringVar(),  #sspValidator
            '$cpfContra' : StringVar(),  #cpfValidator
            '$estadoCivilContra' : StringVar(), #strValidator
            "$compleEmp" : StringVar(), #strValidator
            "$compleContra" : StringVar(), #strValidator
            "$dtVenc" : StringVar(),  #dateValidator
            "$valPag" : StringVar(), #intValidator
            "$dtInic" : StringVar()  #dateValidator
        }
        
    def index(self):
        self.frame_enterprise = Frame(self.window, bd=4, bg='lightblue')
        self.frame_enterprise.place(relx=0.05,rely=0.05,relwidth=0.9,relheight=0.9)

        #Titulo
        Label(self.frame_enterprise, text= self.titulo, background='lightblue', font=('arial',17,'bold'))\
            .place(relx=0.3,rely=0.045)
            
        #Logo
        self.logo = PhotoImage(file='./code/imgs/deltaprice_logo-slim.png')
        
        self.logo = self.logo.subsample(5,5)
        
        Label(self.frame_enterprise, image=self.logo, background='lightblue')\
            .place(relx=0.75,rely=0.01,relwidth=0.12,relheight=0.15)

        #Botão voltar
        Button(self.frame_enterprise, text='Voltar ao menu',\
            command= lambda: (self.frame_enterprise.destroy, self.pageMenu()))\
                .place(relx=0,rely=0,relwidth=0.25,relheight=0.06)

        

        #Labels e Entrys
        #Empresa
        Label(self.frame_enterprise, text='Empresa',\
            background='lightblue', font=('Times New Roman',15,'bold italic'))\
                .place(relx=0.05,rely=0.115)
                
        self.canvas = Canvas(self.frame_enterprise, width=625, height=10, background='darkblue',border=-5)
        self.canvas.place(relx=0.17,rely=0.15)
                
        self.canvas.create_line(-5,0,625,0, fill="darkblue", width=10)
        
        # x,angulo x , y, angulo y
                
        ###########nome empresa

        Label(self.frame_enterprise, text='Nome empresa',\
            background='lightblue', font=(10))\
                .place(relx=0.05,rely=0.18)

        self.nomeEmpEntry = Entry(self.frame_enterprise,\
            textvariable=self.referencias['$nomeEmp'],\
                validate='key', validatecommand=(self.frame_enterprise.register(lambda text: not text.isdecimal()), '%S'))\
                .place(relx=0.05,rely=0.23,relwidth=0.25,relheight=0.05)

        ###########rua

        Label(self.frame_enterprise, text='Rua',\
            background='lightblue', font=(10))\
                .place(relx=0.35,rely=0.18)

        self.ruaEmpEntry = Entry(self.frame_enterprise,\
            textvariable=self.referencias['$ruaEmp'],\
                validate='key', validatecommand=(self.frame_enterprise.register(lambda text: not text.isdecimal()), '%S'))\
                .place(relx=0.35,rely=0.23,relwidth=0.20,relheight=0.05)

        ###########Num

        Label(self.frame_enterprise, text='Num.',\
            background='lightblue', font=(10))\
                .place(relx=0.6,rely=0.18)

        self.numEmpEntry = Entry(self.frame_enterprise,\
            textvariable=self.referencias['$numEmp'],\
                validate='key', validatecommand=(self.frame_enterprise.register(lambda text: text.isdecimal()), '%S'))\
                    .place(relx=0.61,rely=0.23,relwidth=0.05,relheight=0.05)

        ###########bairro

        Label(self.frame_enterprise, text='Bairro',\
            background='lightblue', font=(10))\
                .place(relx=0.7,rely=0.18)

        self.bairroEmpEntry = Entry(self.frame_enterprise,\
            textvariable=self.referencias['$bairroEmp'],\
                validate='key', validatecommand=(self.frame_enterprise.register(lambda text: not text.isdecimal()), '%S'))\
                .place(relx=0.7,rely=0.23,relwidth=0.25,relheight=0.05)

        ########### CEP Empre

        self.valCEP_Empre = StringVar()

        self.valCEP_Empre.trace_add('write', lambda *args, passed = self.valCEP_Empre:\
            Formater.cep_formater(passed, *args) )

        Label(self.frame_enterprise, text='CEP',\
            background='lightblue', font=(10))\
                .place(relx=0.05,rely=0.31)
        

        self.CEPEntry = Entry(self.frame_enterprise, textvariable = self.valCEP_Empre, \
            validate ='key', validatecommand =(self.frame_enterprise.register(Validator.cep_validator), '%P'))\
                .place(relx=0.05,rely=0.36,relwidth=0.25,relheight=0.05)

        self.referencias['$cepEmp'] = self.valCEP_Empre

        ###########TODO CNPJ
        
        self.valCNPJ = StringVar()

        self.valCNPJ.trace_add('write', lambda *args, passed = self.valCNPJ:\
            Formater.cnpj_formater(passed, *args) )

        Label(self.frame_enterprise, text='CNPJ',\
            background='lightblue', font=(10))\
                .place(relx=0.35,rely=0.31)
        

        self.CEPEntry = Entry(self.frame_enterprise, textvariable = self.valCNPJ, \
            validate ='key', validatecommand =(self.frame_enterprise.register(Validator.cnpj_validator), '%P'))\
                .place(relx=0.35,rely=0.36,relwidth=0.2,relheight=0.05)

        self.referencias['$cnpjEmp'] = self.valCNPJ
                
        ###########Complemento

        Label(self.frame_enterprise, text='Complemento',\
            background='lightblue', font=(10))\
                .place(relx=0.6,rely=0.31)

        self.complementoEntry = Entry(self.frame_enterprise,\
            textvariable=self.referencias['$compleEmp'])\
                .place(relx=0.61,rely=0.36,relwidth=0.35,relheight=0.05)
        
        #Socio
        Label(self.frame_enterprise, text='Sócio',\
            background='lightblue', font=('Times New Roman',15,'bold italic'))\
                .place(relx=0.05,rely=0.42)
                
        self.canvas = Canvas(self.frame_enterprise, width=655, height=10,border=-5)
        self.canvas.place(relx=0.13,rely=0.455)
                
        self.canvas.create_line(-5,0,655,0, fill="darkblue", width=10)
        
        # x,angulo x , y, angulo y

        ###########nome

        Label(self.frame_enterprise, text='Nome sócio',\
            background='lightblue', font=(10))\
                .place(relx=0.05,rely=0.48)

        self.nomeEntry = Entry(self.frame_enterprise,\
            textvariable=self.referencias['$nomeContra'],\
                validate='key', validatecommand=(self.frame_enterprise.register(lambda text: not text.isdecimal()), '%S'))\
                .place(relx=0.05,rely=0.53,relwidth=0.25,relheight=0.05)

        ###########rua

        Label(self.frame_enterprise, text='Rua',\
            background='lightblue', font=(10))\
                .place(relx=0.35,rely=0.48)

        self.ruaEntry = Entry(self.frame_enterprise,\
            textvariable=self.referencias['$ruaContra'],\
                validate='key', validatecommand=(self.frame_enterprise.register(lambda text: not text.isdecimal()), '%S'))\
                .place(relx=0.35,rely=0.53,relwidth=0.20,relheight=0.05)

        ###########Num

        Label(self.frame_enterprise, text='Num.',\
            background='lightblue', font=(10))\
                .place(relx=0.6,rely=0.48)

        self.numEntry = Entry(self.frame_enterprise,\
            textvariable=self.referencias['$numContra'],\
                validate='key', validatecommand=(self.frame_enterprise.register(lambda text: text.isdecimal()), '%S'))\
                .place(relx=0.61,rely=0.53,relwidth=0.05,relheight=0.05)

        ###########bairro

        Label(self.frame_enterprise, text='Bairro',\
            background='lightblue', font=(10))\
                .place(relx=0.7,rely=0.48)

        self.bairroEntry = Entry(self.frame_enterprise,\
            textvariable=self.referencias['$bairroContra'],\
                validate='key', validatecommand=(self.frame_enterprise.register(lambda text: not text.isdecimal()), '%S'))\
                .place(relx=0.7,rely=0.53,relwidth=0.25,relheight=0.05)

        ###########TODO CEP
        
        self.valCEP_Contra = StringVar()

        self.valCEP_Contra.trace_add('write', lambda *args, passed = self.valCEP_Contra:\
            Formater.cep_formater(passed, *args) )

        Label(self.frame_enterprise, text='CEP',\
            background='lightblue', font=(10))\
                .place(relx=0.05,rely=0.61)
        

        self.CEPEntry = Entry(self.frame_enterprise, textvariable = self.valCEP_Contra, \
            validate ='key', validatecommand =(self.frame_enterprise.register(Validator.cep_validator), '%P'))\
                .place(relx=0.05,rely=0.66,relwidth=0.25,relheight=0.05)

        self.referencias['$cepContra'] = self.valCEP_Contra

        ###########TODO RG
        
        self.valRG = StringVar()

        self.valRG.trace_add('write', lambda *args, passed = self.valRG:\
            Formater.rg_formater(passed, *args) )

        Label(self.frame_enterprise, text='RG',\
            background='lightblue', font=(10))\
                .place(relx=0.35,rely=0.61)
        

        self.CEPEntry = Entry(self.frame_enterprise, textvariable = self.valRG, \
            validate ='key', validatecommand =(self.frame_enterprise.register(Validator.rg_validator), '%P'))\
                .place(relx=0.35,rely=0.66,relwidth=0.2,relheight=0.05)

        self.referencias['$rgContra'] = self.valRG

        ###########Org. Emissor

        Label(self.frame_enterprise, text='Org.',\
            background='lightblue', font=(10))\
                .place(relx=0.9,rely=0.61)

        self.sspEntry = Entry(self.frame_enterprise,\
            textvariable=self.referencias['$emissorContra'],\
                validate='key', validatecommand=(self.frame_enterprise.register(lambda text: not text.isdecimal()), '%S'))\
                .place(relx=0.9,rely=0.66,relwidth=0.05,relheight=0.05)

        ###########TODO CPF

        self.valCPF = StringVar()

        self.valCPF.trace_add('write', lambda *args, passed = self.valCPF:\
            Formater.cpf_formater(passed, *args) )

        Label(self.frame_enterprise, text='CPF',\
            background='lightblue', font=(10))\
                .place(relx=0.6,rely=0.61)
        

        Entry(self.frame_enterprise, textvariable = self.valCPF, \
            validate ='key', validatecommand =(self.frame_enterprise.register(Validator.cpf_validator), '%P'))\
                .place(relx=0.61,rely=0.66,relwidth=0.25,relheight=0.05)

        self.referencias['$cpfContra'] = self.valCPF

        ###########Estado Civil

        Label(self.frame_enterprise, text='Estado Civil',\
            background='lightblue', font=(10))\
                .place(relx=0.35,rely=0.72)

        self.estadoEntry = StringVar(self.frame_enterprise)

        self.estadoEntryOpt = ('solteiro(a)', 'casado(a)','divorsiado(a)','viuvo(a)')

        self.estadoEntry.set('solteiro(a)')

        self.popup = OptionMenu(self.frame_enterprise, self.estadoEntry, *self.estadoEntryOpt)\
            .place(relx=0.35,rely=0.77,relwidth=0.2,relheight=0.06)

        self.referencias['$estadoCivilContra'] = self.estadoEntry

        ###########Complemento

        Label(self.frame_enterprise, text='Complemento',\
            background='lightblue', font=(10))\
                .place(relx=0.6,rely=0.73)

        self.complementoEntry = Entry(self.frame_enterprise,\
            textvariable=self.referencias['$compleContra'])\
                .place(relx=0.61,rely=0.78,relwidth=0.35,relheight=0.05)

        #Contrato
        Label(self.frame_enterprise, text='Contrato',\
            background='lightblue', font=('Times New Roman',15,'bold italic'))\
                .place(relx=0.05,rely=0.81)
                
        self.canvas = Canvas(self.frame_enterprise, width=625, height=10,border=-5)
        self.canvas.place(relx=0.17,rely=0.845)
                
        self.canvas.create_line(-5,0,625,0, fill="darkblue", width=10)
        
        # x,angulo x , y, angulo y

        ###########TODO Valor pagamento

        self.valPag = StringVar(value='R$ ')

        Label(self.frame_enterprise, text='Val. Pagamento',\
            background='lightblue', font=(10))\
                .place(relx=0.05,rely=0.88)
        
        self.valEntry = Entry(self.frame_enterprise, textvariable = self.valPag, )\
                .place(relx=0.05,rely=0.93,relwidth=0.1,relheight=0.05)

        self.referencias['$valPag'] = self.valPag

        ###########TODO Data inicio
        
        self.valDT_inic = StringVar()

        self.valDT_inic.trace_add('write', lambda *args, passed = self.valDT_inic:\
            Formater.date_formater(passed, *args) )

        Label(self.frame_enterprise, text='Data início',\
            background='lightblue', font=(10))\
                .place(relx=0.25,rely=0.88)
        

        Entry(self.frame_enterprise, textvariable = self.valDT_inic, \
            validate ='key', validatecommand =(self.frame_enterprise.register(Validator.date_validator), '%P'))\
                .place(relx=0.25,rely=0.93,relwidth=0.1,relheight=0.05)

        self.referencias['$dtInic'] = self.valDT_inic

        ###########TODO Data vencimento
        
        self.valDT_venc = StringVar()

        self.valDT_venc.trace_add('write', lambda *args, passed = self.valDT_venc:\
            Formater.date_formater(passed, *args) )

        Label(self.frame_enterprise, text='Data vencimento',\
            background='lightblue', font=(10))\
                .place(relx=0.4,rely=0.88)
        

        Entry(self.frame_enterprise, textvariable = self.valDT_venc, \
            validate ='key', validatecommand =(self.frame_enterprise.register(Validator.date_validator), '%P'))\
                .place(relx=0.4,rely=0.93,relwidth=0.1,relheight=0.05)

        self.referencias['$dtVenc'] = self.valDT_venc

        #Botão enviar
        Button(self.frame_enterprise, text='Gerar CPS',\
            command= lambda: self.alterar_doc(frame_ativo=self.frame_enterprise))\
                .place(relx=0.61,rely=0.865,relwidth=0.35,relheight=0.12)
        
    def changes(self):
        estadoCiv = self.referencias['$estadoCivilContra']
        if 'STB' in estadoCiv.get():
            estadoCiv.set('Casado em Separação Total de Bens')
        elif 'CPB' in estadoCiv.get():
            estadoCiv.set('Casado em Comunhão Parcial de Bens')
        elif 'CTB' in estadoCiv.get():
            estadoCiv.set('Casado em Comunhão Total de Bens')
            
        # valor = self.referencias['$valPag']
        # valorDbl = valor.get()[2:].replace(',','.')
        # valorExtenso = num2words(valorDbl,lang='pt-br')
        # valor.set(f'{valor.get()} ({valorExtenso})') 
    
class Person(Pages):
    def __init__(self):
        super().__init__(titulo, caminho)
        self.referencias = {
            '$nomeContra' : StringVar(), #strValidator
            '$ruaContra' : StringVar(), #strValidator
            '$numContra' : StringVar(), #intValidator
            '$bairroContra' : StringVar(),  #strValidator
            '$cepContra' : StringVar(),  #intValidator
            '$rgContra' : StringVar(),  #rgValidator
            '$emissorContra' : StringVar(),  #sspValidator
            '$cpfContra' : StringVar(),  #cpfValidator
            '$estadoCivilContra' : StringVar(), #strValidator
            "$compleEmp" : StringVar(), #strValidator
            "$compleContra" : StringVar(), #strValidator
            "$dtVenc" : StringVar(),  #dateValidator
            "$valPag" : StringVar(), #intValidator
            "$dtInic" : StringVar()  #dateValidator
        }
    
    def index(self):
        self.frame_Person = Frame(self.window, bd=4, bg='lightblue')
        self.frame_Person.place(relx=0.05,rely=0.05,relwidth=0.9,relheight=0.9)
        

        #Titulo
        Label(self.frame_Person, text='Gerador de CPS Pessoa física', background='lightblue', font=('arial',17,'bold'))\
            .place(relx=0.3,rely=0.05)

        #Logo
        self.logo = PhotoImage(file='./code/imgs/deltaprice_logo-slim.png')
        
        self.logo = self.logo.subsample(5,5)
        
        Label(self.frame_Person, image=self.logo, background='lightblue')\
            .place(relx=0.75,rely=0.09,relwidth=0.12,relheight=0.15)

        #Botão voltar
        Button(self.frame_Person, text='Voltar ao menu',\
            command= lambda: (self.frame_Person.destroy, self.pageMenu()))\
                .place(relx=0,rely=0,relwidth=0.25,relheight=0.06)

        self.referencias = {
            '$nomeContra' : StringVar(), #strValidator
            '$rgContra' : StringVar(),  #rgValidator
            '$emissorContra' : StringVar(),  #sspValidator
            '$cpfContra' : StringVar(),  #cpfValidator
            '$estadoCivilContra' : StringVar(), #strValidator
            '$estadoContra' : StringVar(), #strValidator
            '$cidadeContra' : StringVar(), #strValidator
            '$ruaContra' : StringVar(), #strValidator
            '$bairroContra' : StringVar(),  #strValidator
            '$numContra' : StringVar(), #intValidator
            "$compleContra" : StringVar(), #strValidator
            '$cepContra' : StringVar(),  #intValidator
            "$numEmpre" : StringVar(),  #dateValidator
            "$valPag" : StringVar(), #valValidator
            "$dtVenc" : StringVar(),  #dateValidator
            "$dtAss" : StringVar(),  #dateValidator
            "$dtInic" : StringVar()  #dateValidator
        }

        #Labels e Entrys
        #Contratante
        Label(self.frame_Person, text='Contratante',\
            background='lightblue', font=('Times New Roman',15,'bold italic'))\
                .place(relx=0.05,rely=0.2)

        ############Linha
        self.canvas = Canvas(self.frame_Person, width=600, height=10, background='darkblue',border=-5)
        self.canvas.place(relx=0.2,rely=0.23)
                
        self.canvas.create_line(-5,0,600,0, fill="darkblue", width=10)

        ###########nome

        Label(self.frame_Person, text='Nome',\
            background='lightblue', font=(10))\
                .place(relx=0.05,rely=0.3)

        Entry(self.frame_Person,\
            textvariable=self.referencias['$nomeContra'],\
                validate='key', validatecommand=(self.frame_Person.register(lambda text: not text.isdecimal()), '%S'))\
                .place(relx=0.05,rely=0.37,relwidth=0.25,relheight=0.05)
        
         ###########RG
        
        self.valRG = StringVar()

        self.valRG.trace_add('write', lambda *args, passed = self.valRG:\
            Formater.rg_formater(passed, *args) )

        Label(self.frame_Person, text='RG',\
            background='lightblue', font=(10))\
                .place(relx=0.33,rely=0.3)
        

        Entry(self.frame_Person, textvariable = self.valRG, \
            validate ='key', validatecommand =(self.frame_Person.register(Validator.rg_validator), '%P'))\
                .place(relx=0.33,rely=0.37,relwidth=0.15,relheight=0.05)

        self.referencias['$rgContra'] = self.valRG
        
        ###########Org. Emissor

        Label(self.frame_Person, text='Org.',\
            background='lightblue', font=(10))\
                .place(relx=0.5,rely=0.3)

        Entry(self.frame_Person,\
            textvariable=self.referencias['$emissorContra'],\
                validate='key', validatecommand=(self.frame_Person.register(lambda text: not text.isdecimal()), '%S'))\
                .place(relx=0.5,rely=0.37,relwidth=0.05,relheight=0.05)

        ###########CPF
        
        self.valCPF = StringVar()

        self.valCPF.trace_add('write', lambda *args, passed = self.valCPF:\
            Formater.cpf_formater(passed, *args) )

        Label(self.frame_Person, text='CPF',\
            background='lightblue', font=(10))\
                .place(relx=0.575,rely=0.3)
        

        Entry(self.frame_Person, textvariable = self.valCPF, \
            validate ='key', validatecommand =(self.frame_Person.register(Validator.cpf_validator), '%P'))\
                .place(relx=0.575,rely=0.37,relwidth=0.15,relheight=0.05)

        self.referencias['$cpfContra'] = self.valCPF

        ###########Estado Civil
        
        Label(self.frame_Person, text='Estado Civil',\
            background='lightblue', font=(10))\
                .place(relx=0.75,rely=0.3)

        self.estadoEntry = StringVar(self.frame_Person)

        self.estadoEntryOpt = ('solteiro(a)','divorsiado(a)','viuvo(a)')

        self.popup = ttk.OptionMenu(self.frame_Person, self.estadoEntry,'', *self.estadoEntryOpt)

        self.menuCasado = self.popup['menu']

        #Casado
        self.subLista = Menu(self.menuCasado, tearoff=False)
        self.menuCasado.add_cascade(label = 'casado(a)',menu= self.subLista)
        self.subLista.add_command(label='Comunhão Parcial de Bens', \
            command= lambda: self.estadoEntry.set('casado(a) em CPB'))
        
        self.subLista.add_command(label='Comunhão Total de Bens',\
            command= lambda: self.estadoEntry.set('casado(a) em CTB'))
        
        self.subLista.add_command(label='Separação Total de Bens',\
            command= lambda: self.estadoEntry.set('casado(a) em STB'))


        self.popup.place(relx=0.75,rely=0.37,relwidth=0.2,relheight=0.06)

        self.referencias['$estadoCivilContra'] = self.estadoEntry

        ###########rua

        Label(self.frame_Person, text='Rua',\
            background='lightblue', font=(10))\
                .place(relx=0.05,rely=0.45)

        Entry(self.frame_Person,\
            textvariable=self.referencias['$ruaContra'],\
                validate='key', validatecommand=(self.frame_Person.register(lambda text: not text.isdecimal()), '%S'))\
                .place(relx=0.05,rely=0.52,relwidth=0.25,relheight=0.05)

        ###########Num

        Label(self.frame_Person, text='Num.',\
            background='lightblue', font=(10))\
                .place(relx=0.33,rely=0.45)

        Entry(self.frame_Person,\
            textvariable=self.referencias['$numContra'],\
                validate='key', validatecommand=(self.frame_Person.register(lambda text: text.isdecimal()), '%S'))\
                .place(relx=0.33,rely=0.52,relwidth=0.05,relheight=0.05)

        ###########bairro

        Label(self.frame_Person, text='Bairro',\
            background='lightblue', font=(10))\
                .place(relx=0.415,rely=0.45)

        Entry(self.frame_Person,\
            textvariable=self.referencias['$bairroContra'],\
                validate='key', validatecommand=(self.frame_Person.register(lambda text: not text.isdecimal()), '%S'))\
                .place(relx=0.415,rely=0.52,relwidth=0.25,relheight=0.05)
        
        ###########CEP 
        
        self.valCEP_Contra = StringVar()

        self.valCEP_Contra.trace_add('write', lambda *args, passed = self.valCEP_Contra:\
            Formater.cep_formater(passed, *args) )

        Label(self.frame_Person, text='CEP',\
            background='lightblue', font=(10))\
                .place(relx=0.7,rely=0.45)
        

        Entry(self.frame_Person, textvariable = self.valCEP_Contra, \
            validate ='key', validatecommand =(self.frame_Person.register(Validator.cep_validator), '%P'))\
                .place(relx=0.7,rely=0.52,relwidth=0.25,relheight=0.05)

        self.referencias['$cepContra'] = self.valCEP_Contra
        
        ###########Cidade

        Label(self.frame_Person, text='Cidade',\
            background='lightblue', font=(10))\
                .place(relx=0.05,rely=0.6)

        Entry(self.frame_Person,\
            textvariable=self.referencias['$cidadeContra'],\
                validate='key', validatecommand=(self.frame_Person.register(lambda text: not text.isdecimal()), '%S'))\
                .place(relx=0.05,rely=0.67,relwidth=0.25,relheight=0.05)
        
        ###########Estado

        Label(self.frame_Person, text='Estado',\
            background='lightblue', font=(10))\
                .place(relx=0.33,rely=0.6)

        Entry(self.frame_Person,\
            textvariable=self.referencias['$estadoContra'],\
                validate='key', validatecommand=(self.frame_Person.register(lambda text: not text.isdecimal()), '%S'))\
                .place(relx=0.33,rely=0.67,relwidth=0.25,relheight=0.05)

        ###########Complemento

        Label(self.frame_Person, text='Complemento (opcional)',\
            background='lightblue', font=(10))\
                .place(relx=0.6,rely=0.6)

        Entry(self.frame_Person,\
            textvariable=self.referencias['$compleContra'])\
                .place(relx=0.61,rely=0.67,relwidth=0.34,relheight=0.05)

         #Contrato
        Label(self.frame_Person, text='Contrato',\
            background='lightblue', font=('Times New Roman',15,'bold italic'))\
                .place(relx=0.05,rely=0.75)

        ############Linha
        self.canvas = Canvas(self.frame_Person, width=625, height=10, background='darkblue',border=-5)
        self.canvas.place(relx=0.17,rely=0.78)
                
        self.canvas.create_line(-5,0,625,0, fill="darkblue", width=10)

        ###########Valor pagamento
        
        self.valPag = StringVar(value='R$ ')

        self.valPag.trace_add('write', lambda *args, passed = self.valPag:\
            Formater.valor_formater(passed, *args) )

        Label(self.frame_Person, text='Val. Contrato.',\
            background='lightblue', font=(10))\
                .place(relx=0.05,rely=0.85)
        
        self.entryVal = Entry(self.frame_Person, textvariable = self.valPag, )\
                .place(relx=0.06,rely=0.92,relwidth=0.1,relheight=0.05)
                
        self.referencias['$valPag'] = self.valPag

        ###########Data inicio
        
        self.valDT_inic = StringVar()

        self.valDT_inic.trace_add('write', lambda *args, passed = self.valDT_inic:\
            Formater.date_formater(passed, *args) )

        Label(self.frame_Person, text='Data início',\
            background='lightblue', font=(10))\
                .place(relx=0.185,rely=0.85)
        

        Entry(self.frame_Person, textvariable = self.valDT_inic, \
            validate ='key', validatecommand =(self.frame_Person.register(Validator.date_validator), '%P')).place(relx=0.185,rely=0.92,relwidth=0.1,relheight=0.05)

        self.referencias['$dtInic'] = self.valDT_inic

        ###########Data Assinatura
        
        self.valDT_ass = StringVar()

        self.valDT_ass.trace_add('write', lambda *args, passed = self.valDT_ass:\
            Formater.date_formater(passed, *args) )

        Label(self.frame_Person, text='Data Ass.',\
            background='lightblue', font=(10))\
                .place(relx=0.3,rely=0.85)
        

        Entry(self.frame_Person, textvariable = self.valDT_ass, \
            validate ='key', validatecommand =(self.frame_Person.register(Validator.date_validator), '%P')).place(relx=0.3,rely=0.92,relwidth=0.1,relheight=0.05)

        self.referencias['$dtAss'] = self.valDT_ass

        ###########Dia vencimento
        
        Label(self.frame_Person, text='Dia Venc.',\
            background='lightblue', font=(10))\
                .place(relx=0.42,rely=0.85)
        

        Entry(self.frame_Person,textvariable=self.referencias['$dtVenc'], validate='key', validatecommand=(self.frame_Person.register(lambda text: text.isdecimal()), '%S')).place(relx=0.42,rely=0.92,relwidth=0.1,relheight=0.05)


        ###########Num. Empregados

        Label(self.frame_Person, text='Num. Empreg.',\
            background='lightblue', font=(10))\
                .place(relx=0.55,rely=0.85)

        Entry(self.frame_Person,\
            textvariable=self.referencias['$numEmpre'],\
                validate='key', validatecommand=(self.frame_Person.register(lambda text: text.isdecimal()), '%S'))\
                .place(relx=0.58,rely=0.92,relwidth=0.05,relheight=0.05)

        #Botão enviar
        Button(self.frame_Person, text='Gerar CPS',\
            command= lambda: self.alterar_doc(frame_ativo=self.frame_Person))\
                .place(relx=0.7,rely=0.85,relwidth=0.25,relheight=0.12)
     
class PessoaFisica(Person):
    def __init__(self):
        super().__init__('Pessoa Fisica', './code/CPS\'s/CPS INATIVIDADE.docx')
    
class Inatividade(Enter):
    def __init__(self):
        super().__init__('Inatividade', './code/CPS\'s/CPS INATIVIDADE.docx')
        
class LucroPresumido(Empresa):
    def __init__(self):
        super().__init__('Inatividade', './code/CPS\'s/CPS INATIVIDADE.docx')
        
class SimplesNacional(Empresa):
    def __init__(self):
        super().__init__('Inatividade', './code/CPS\'s/CPS INATIVIDADE.docx')

class App:
    def __init__(self):
        self.window = window
        self.pages = Pages()
        self.tela()
        self.pageMenu()
        window.mainloop()

    def tela(self):
        self.window.configure(background='darkblue')
        self.window.resizable(False,False)
        self.window.geometry('880x500')
        self.window.iconbitmap('./code/imgs/delta-icon.ico')
        self.window.title('Gerador de CPS')

    def menu(self):
        self.menu = Frame(self.window, bd=4, bg='lightblue')
        self.menu.place(relx=0.05,rely=0.05,relwidth=0.9,relheight=0.9)

        self.textOrientacao = Label(self.menu, text='Selecione o tipo de CPS que deseja fazer:', background='lightblue', font=('arial',20,'bold'))\
        .place(relx=0.15,rely=0.2,relheight=0.15)
        
        #Logo
        self.logo = PhotoImage(file='./code/imgs/deltaprice-hori.png')
        
        self.logo = self.logo.subsample(4,4)
        
        Label(self.window, image=self.logo, background='lightblue')\
            .place(relx=0.175,rely=0.05,relwidth=0.7,relheight=0.2)

        #Pessoa física
        self.btnPF = Button(self.menu, text='CPS Pessoa Física',\
            command= lambda: PessoaFisica.index())\
                .place(relx=0.15,rely=0.7,relwidth=0.25,relheight=0.15)

        #Inatividade
        self.btnIN = Button(self.menu, text='CPS Inatividade',\
            command= lambda: Inatividade.index())\
                .place(relx=0.60,rely=0.4,relwidth=0.25,relheight=0.15)

        #Lucro Presumido
        self.btnLP = Button(self.menu, text='CPS Lucros',\
            command= lambda: LucroPresumido.index())\
                .place(relx=0.15,rely=0.4,relwidth=0.25,relheight=0.15)

        #Simples Nacional
        self.btnSN = Button(self.menu, text='CPS Simples Nacional',\
            command= lambda: SimplesNacional.index())\
                .place(relx=0.60,rely=0.7,relwidth=0.25,relheight=0.15)

    def pagePF(self):
        self.menu.destroy()
        self.doc = Document('./code/CPS\'s/CPS PESSOA FISICA.docx')

        self.cpsPF = Frame(self.window, bd=4, bg='lightblue')
        self.cpsPF.place(relx=0.05,rely=0.05,relwidth=0.9,relheight=0.9)
        

        #Titulo
        Label(self.cpsPF, text='Gerador de CPS Pessoa física', background='lightblue', font=('arial',17,'bold'))\
            .place(relx=0.3,rely=0.05)

        #Logo
        self.logo = PhotoImage(file='./code/imgs/deltaprice_logo-slim.png')
        
        self.logo = self.logo.subsample(5,5)
        
        Label(self.cpsPF, image=self.logo, background='lightblue')\
            .place(relx=0.75,rely=0.09,relwidth=0.12,relheight=0.15)

        #Botão voltar
        Button(self.cpsPF, text='Voltar ao menu',\
            command= lambda: (self.cpsPF.destroy, self.pageMenu()))\
                .place(relx=0,rely=0,relwidth=0.25,relheight=0.06)

        self.referencias = {
            '$nomeContra' : StringVar(), #strValidator
            '$rgContra' : StringVar(),  #rgValidator
            '$emissorContra' : StringVar(),  #sspValidator
            '$cpfContra' : StringVar(),  #cpfValidator
            '$estadoCivilContra' : StringVar(), #strValidator
            '$estadoContra' : StringVar(), #strValidator
            '$cidadeContra' : StringVar(), #strValidator
            '$ruaContra' : StringVar(), #strValidator
            '$bairroContra' : StringVar(),  #strValidator
            '$numContra' : StringVar(), #intValidator
            "$compleContra" : StringVar(), #strValidator
            '$cepContra' : StringVar(),  #intValidator
            "$numEmpre" : StringVar(),  #dateValidator
            "$valPag" : StringVar(), #valValidator
            "$dtVenc" : StringVar(),  #dateValidator
            "$dtAss" : StringVar(),  #dateValidator
            "$dtInic" : StringVar()  #dateValidator
        }

        #Labels e Entrys
        #Contratante
        Label(self.cpsPF, text='Contratante',\
            background='lightblue', font=('Times New Roman',15,'bold italic'))\
                .place(relx=0.05,rely=0.2)

        ############Linha
        self.canvas = Canvas(self.cpsPF, width=600, height=10, background='darkblue',border=-5)
        self.canvas.place(relx=0.2,rely=0.23)
                
        self.canvas.create_line(-5,0,600,0, fill="darkblue", width=10)

        ###########nome

        Label(self.cpsPF, text='Nome',\
            background='lightblue', font=(10))\
                .place(relx=0.05,rely=0.3)

        Entry(self.cpsPF,\
            textvariable=self.referencias['$nomeContra'],\
                validate='key', validatecommand=(self.cpsPF.register(lambda text: not text.isdecimal()), '%S'))\
                .place(relx=0.05,rely=0.37,relwidth=0.25,relheight=0.05)
        
         ###########RG
        
        self.valRG = StringVar()

        self.valRG.trace_add('write', lambda *args, passed = self.valRG:\
            Formater.rg_formater(passed, *args) )

        Label(self.cpsPF, text='RG',\
            background='lightblue', font=(10))\
                .place(relx=0.33,rely=0.3)
        

        Entry(self.cpsPF, textvariable = self.valRG, \
            validate ='key', validatecommand =(self.cpsPF.register(Validator.rg_validator), '%P'))\
                .place(relx=0.33,rely=0.37,relwidth=0.15,relheight=0.05)

        self.referencias['$rgContra'] = self.valRG
        
        ###########Org. Emissor

        Label(self.cpsPF, text='Org.',\
            background='lightblue', font=(10))\
                .place(relx=0.5,rely=0.3)

        Entry(self.cpsPF,\
            textvariable=self.referencias['$emissorContra'],\
                validate='key', validatecommand=(self.cpsPF.register(lambda text: not text.isdecimal()), '%S'))\
                .place(relx=0.5,rely=0.37,relwidth=0.05,relheight=0.05)

        ###########CPF
        
        self.valCPF = StringVar()

        self.valCPF.trace_add('write', lambda *args, passed = self.valCPF:\
            Formater.cpf_formater(passed, *args) )

        Label(self.cpsPF, text='CPF',\
            background='lightblue', font=(10))\
                .place(relx=0.575,rely=0.3)
        

        Entry(self.cpsPF, textvariable = self.valCPF, \
            validate ='key', validatecommand =(self.cpsPF.register(Validator.cpf_validator), '%P'))\
                .place(relx=0.575,rely=0.37,relwidth=0.15,relheight=0.05)

        self.referencias['$cpfContra'] = self.valCPF

        ###########Estado Civil
        
        Label(self.cpsPF, text='Estado Civil',\
            background='lightblue', font=(10))\
                .place(relx=0.75,rely=0.3)

        self.estadoEntry = StringVar(self.cpsPF)

        self.estadoEntryOpt = ('solteiro(a)','divorsiado(a)','viuvo(a)')

        self.popup = ttk.OptionMenu(self.cpsPF, self.estadoEntry,'', *self.estadoEntryOpt)

        self.menuCasado = self.popup['menu']

        #Casado
        self.subLista = Menu(self.menuCasado, tearoff=False)
        self.menuCasado.add_cascade(label = 'casado(a)',menu= self.subLista)
        self.subLista.add_command(label='Comunhão Parcial de Bens', \
            command= lambda: self.estadoEntry.set('casado(a) em CPB'))
        
        self.subLista.add_command(label='Comunhão Total de Bens',\
            command= lambda: self.estadoEntry.set('casado(a) em CTB'))
        
        self.subLista.add_command(label='Separação Total de Bens',\
            command= lambda: self.estadoEntry.set('casado(a) em STB'))


        self.popup.place(relx=0.75,rely=0.37,relwidth=0.2,relheight=0.06)

        self.referencias['$estadoCivilContra'] = self.estadoEntry

        ###########rua

        Label(self.cpsPF, text='Rua',\
            background='lightblue', font=(10))\
                .place(relx=0.05,rely=0.45)

        Entry(self.cpsPF,\
            textvariable=self.referencias['$ruaContra'],\
                validate='key', validatecommand=(self.cpsPF.register(lambda text: not text.isdecimal()), '%S'))\
                .place(relx=0.05,rely=0.52,relwidth=0.25,relheight=0.05)

        ###########Num

        Label(self.cpsPF, text='Num.',\
            background='lightblue', font=(10))\
                .place(relx=0.33,rely=0.45)

        Entry(self.cpsPF,\
            textvariable=self.referencias['$numContra'],\
                validate='key', validatecommand=(self.cpsPF.register(lambda text: text.isdecimal()), '%S'))\
                .place(relx=0.33,rely=0.52,relwidth=0.05,relheight=0.05)

        ###########bairro

        Label(self.cpsPF, text='Bairro',\
            background='lightblue', font=(10))\
                .place(relx=0.415,rely=0.45)

        Entry(self.cpsPF,\
            textvariable=self.referencias['$bairroContra'],\
                validate='key', validatecommand=(self.cpsPF.register(lambda text: not text.isdecimal()), '%S'))\
                .place(relx=0.415,rely=0.52,relwidth=0.25,relheight=0.05)
        
        ###########CEP 
        
        self.valCEP_Contra = StringVar()

        self.valCEP_Contra.trace_add('write', lambda *args, passed = self.valCEP_Contra:\
            Formater.cep_formater(passed, *args) )

        Label(self.cpsPF, text='CEP',\
            background='lightblue', font=(10))\
                .place(relx=0.7,rely=0.45)
        

        Entry(self.cpsPF, textvariable = self.valCEP_Contra, \
            validate ='key', validatecommand =(self.cpsPF.register(Validator.cep_validator), '%P'))\
                .place(relx=0.7,rely=0.52,relwidth=0.25,relheight=0.05)

        self.referencias['$cepContra'] = self.valCEP_Contra
        
        ###########Cidade

        Label(self.cpsPF, text='Cidade',\
            background='lightblue', font=(10))\
                .place(relx=0.05,rely=0.6)

        Entry(self.cpsPF,\
            textvariable=self.referencias['$cidadeContra'],\
                validate='key', validatecommand=(self.cpsPF.register(lambda text: not text.isdecimal()), '%S'))\
                .place(relx=0.05,rely=0.67,relwidth=0.25,relheight=0.05)
        
        ###########Estado

        Label(self.cpsPF, text='Estado',\
            background='lightblue', font=(10))\
                .place(relx=0.33,rely=0.6)

        Entry(self.cpsPF,\
            textvariable=self.referencias['$estadoContra'],\
                validate='key', validatecommand=(self.cpsPF.register(lambda text: not text.isdecimal()), '%S'))\
                .place(relx=0.33,rely=0.67,relwidth=0.25,relheight=0.05)

        ###########Complemento

        Label(self.cpsPF, text='Complemento (opcional)',\
            background='lightblue', font=(10))\
                .place(relx=0.6,rely=0.6)

        Entry(self.cpsPF,\
            textvariable=self.referencias['$compleContra'])\
                .place(relx=0.61,rely=0.67,relwidth=0.34,relheight=0.05)

         #Contrato
        Label(self.cpsPF, text='Contrato',\
            background='lightblue', font=('Times New Roman',15,'bold italic'))\
                .place(relx=0.05,rely=0.75)

        ############Linha
        self.canvas = Canvas(self.cpsPF, width=625, height=10, background='darkblue',border=-5)
        self.canvas.place(relx=0.17,rely=0.78)
                
        self.canvas.create_line(-5,0,625,0, fill="darkblue", width=10)

        ###########Valor pagamento
        
        self.valPag = StringVar(value='R$ ')

        self.valPag.trace_add('write', lambda *args, passed = self.valPag:\
            Formater.valor_formater(passed, *args) )

        Label(self.cpsPF, text='Val. Contrato.',\
            background='lightblue', font=(10))\
                .place(relx=0.05,rely=0.85)
        
        self.entryVal = Entry(self.cpsPF, textvariable = self.valPag, )\
                .place(relx=0.06,rely=0.92,relwidth=0.1,relheight=0.05)
                
        self.referencias['$valPag'] = self.valPag

        ###########Data inicio
        
        self.valDT_inic = StringVar()

        self.valDT_inic.trace_add('write', lambda *args, passed = self.valDT_inic:\
            Formater.date_formater(passed, *args) )

        Label(self.cpsPF, text='Data início',\
            background='lightblue', font=(10))\
                .place(relx=0.185,rely=0.85)
        

        Entry(self.cpsPF, textvariable = self.valDT_inic, \
            validate ='key', validatecommand =(self.cpsPF.register(Validator.date_validator), '%P')).place(relx=0.185,rely=0.92,relwidth=0.1,relheight=0.05)

        self.referencias['$dtInic'] = self.valDT_inic

        ###########Data Assinatura
        
        self.valDT_ass = StringVar()

        self.valDT_ass.trace_add('write', lambda *args, passed = self.valDT_ass:\
            Formater.date_formater(passed, *args) )

        Label(self.cpsPF, text='Data Ass.',\
            background='lightblue', font=(10))\
                .place(relx=0.3,rely=0.85)
        

        Entry(self.cpsPF, textvariable = self.valDT_ass, \
            validate ='key', validatecommand =(self.cpsPF.register(Validator.date_validator), '%P')).place(relx=0.3,rely=0.92,relwidth=0.1,relheight=0.05)

        self.referencias['$dtAss'] = self.valDT_ass

        ###########Dia vencimento
        
        Label(self.cpsPF, text='Dia Venc.',\
            background='lightblue', font=(10))\
                .place(relx=0.42,rely=0.85)
        

        Entry(self.cpsPF,textvariable=self.referencias['$dtVenc'], validate='key', validatecommand=(self.cpsPF.register(lambda text: text.isdecimal()), '%S')).place(relx=0.42,rely=0.92,relwidth=0.1,relheight=0.05)


        ###########Num. Empregados

        Label(self.cpsPF, text='Num. Empreg.',\
            background='lightblue', font=(10))\
                .place(relx=0.55,rely=0.85)

        Entry(self.cpsPF,\
            textvariable=self.referencias['$numEmpre'],\
                validate='key', validatecommand=(self.cpsPF.register(lambda text: text.isdecimal()), '%S'))\
                .place(relx=0.58,rely=0.92,relwidth=0.05,relheight=0.05)

        #Botão enviar
        Button(self.cpsPF, text='Gerar CPS',\
            command= lambda: self.alterar_doc(frame_ativo=self.cpsPF))\
                .place(relx=0.7,rely=0.85,relwidth=0.25,relheight=0.12)
        
    def pageIN(self):
        self.menu.destroy()
        self.doc = Document('./code/CPS\'s/CPS INATIVIDADE.docx')

        self.cpsIN = Frame(self.window, bd=4, bg='lightblue')
        self.cpsIN.place(relx=0.05,rely=0.05,relwidth=0.9,relheight=0.9)

        #Titulo
        Label(self.cpsIN, text='Gerador de CPS Inatividade', background='lightblue', font=('arial',17,'bold'))\
            .place(relx=0.3,rely=0.045)
            
        #Logo
        self.logo = PhotoImage(file='./code/imgs/deltaprice_logo-slim.png')
        
        self.logo = self.logo.subsample(5,5)
        
        Label(self.cpsIN, image=self.logo, background='lightblue')\
            .place(relx=0.75,rely=0.01,relwidth=0.12,relheight=0.15)

        #Botão voltar
        Button(self.cpsIN, text='Voltar ao menu',\
            command= lambda: (self.cpsIN.destroy, self.pageMenu()))\
                .place(relx=0,rely=0,relwidth=0.25,relheight=0.06)

        self.referencias = {
            '$nomeEmp' : StringVar(), #strValidator
            '$ruaEmp' : StringVar(), #strValidator
            '$numEmp' : StringVar(), #intValidator
            '$bairroEmp' : StringVar(),  #strValidator
            '$cepEmp' : StringVar(),  #intValidator
            '$rgEmp' : StringVar(),  #rgValidator
            '$sspEmp' : StringVar(),  #sspValidator
            '$cnpjEmp' : StringVar(),  #cpfValidator
            '$nomeContra' : StringVar(), #strValidator
            '$ruaContra' : StringVar(), #strValidator
            '$numContra' : StringVar(), #intValidator
            '$bairroContra' : StringVar(),  #strValidator
            '$cepContra' : StringVar(),  #intValidator
            '$rgContra' : StringVar(),  #rgValidator
            '$emissorContra' : StringVar(),  #sspValidator
            '$cpfContra' : StringVar(),  #cpfValidator
            '$estadoCivilContra' : StringVar(), #strValidator
            "$compleEmp" : StringVar(), #strValidator
            "$compleContra" : StringVar(), #strValidator
            "$dtVenc" : StringVar(),  #dateValidator
            "$valPag" : StringVar(), #intValidator
            "$dtInic" : StringVar()  #dateValidator
        }

        #Labels e Entrys
        #Empresa
        Label(self.cpsIN, text='Empresa',\
            background='lightblue', font=('Times New Roman',15,'bold italic'))\
                .place(relx=0.05,rely=0.115)
                
        self.canvas = Canvas(self.cpsIN, width=625, height=10, background='darkblue',border=-5)
        self.canvas.place(relx=0.17,rely=0.15)
                
        self.canvas.create_line(-5,0,625,0, fill="darkblue", width=10)
        
        # x,angulo x , y, angulo y
                
        ###########nome empresa

        Label(self.cpsIN, text='Nome empresa',\
            background='lightblue', font=(10))\
                .place(relx=0.05,rely=0.18)

        self.nomeEmpEntry = Entry(self.cpsIN,\
            textvariable=self.referencias['$nomeEmp'],\
                validate='key', validatecommand=(self.cpsIN.register(lambda text: not text.isdecimal()), '%S'))\
                .place(relx=0.05,rely=0.23,relwidth=0.25,relheight=0.05)

        ###########rua

        Label(self.cpsIN, text='Rua',\
            background='lightblue', font=(10))\
                .place(relx=0.35,rely=0.18)

        self.ruaEmpEntry = Entry(self.cpsIN,\
            textvariable=self.referencias['$ruaEmp'],\
                validate='key', validatecommand=(self.cpsIN.register(lambda text: not text.isdecimal()), '%S'))\
                .place(relx=0.35,rely=0.23,relwidth=0.20,relheight=0.05)

        ###########Num

        Label(self.cpsIN, text='Num.',\
            background='lightblue', font=(10))\
                .place(relx=0.6,rely=0.18)

        self.numEmpEntry = Entry(self.cpsIN,\
            textvariable=self.referencias['$numEmp'],\
                validate='key', validatecommand=(self.cpsIN.register(lambda text: text.isdecimal()), '%S'))\
                    .place(relx=0.61,rely=0.23,relwidth=0.05,relheight=0.05)

        ###########bairro

        Label(self.cpsIN, text='Bairro',\
            background='lightblue', font=(10))\
                .place(relx=0.7,rely=0.18)

        self.bairroEmpEntry = Entry(self.cpsIN,\
            textvariable=self.referencias['$bairroEmp'],\
                validate='key', validatecommand=(self.cpsIN.register(lambda text: not text.isdecimal()), '%S'))\
                .place(relx=0.7,rely=0.23,relwidth=0.25,relheight=0.05)

        ########### CEP Empre

        self.valCEP_Empre = StringVar()

        self.valCEP_Empre.trace_add('write', lambda *args, passed = self.valCEP_Empre:\
            Formater.cep_formater(passed, *args) )

        Label(self.cpsIN, text='CEP',\
            background='lightblue', font=(10))\
                .place(relx=0.05,rely=0.31)
        

        self.CEPEntry = Entry(self.cpsIN, textvariable = self.valCEP_Empre, \
            validate ='key', validatecommand =(self.cpsIN.register(Validator.cep_validator), '%P'))\
                .place(relx=0.05,rely=0.36,relwidth=0.25,relheight=0.05)

        self.referencias['$cepEmp'] = self.valCEP_Empre

        ###########TODO CNPJ
        
        self.valCNPJ = StringVar()

        self.valCNPJ.trace_add('write', lambda *args, passed = self.valCNPJ:\
            Formater.cnpj_formater(passed, *args) )

        Label(self.cpsIN, text='CNPJ',\
            background='lightblue', font=(10))\
                .place(relx=0.35,rely=0.31)
        

        self.CEPEntry = Entry(self.cpsIN, textvariable = self.valCNPJ, \
            validate ='key', validatecommand =(self.cpsIN.register(Validator.cnpj_validator), '%P'))\
                .place(relx=0.35,rely=0.36,relwidth=0.2,relheight=0.05)

        self.referencias['$cnpjEmp'] = self.valCNPJ
                
        ###########Complemento

        Label(self.cpsIN, text='Complemento',\
            background='lightblue', font=(10))\
                .place(relx=0.6,rely=0.31)

        self.complementoEntry = Entry(self.cpsIN,\
            textvariable=self.referencias['$compleEmp'])\
                .place(relx=0.61,rely=0.36,relwidth=0.35,relheight=0.05)
        
        #Socio
        Label(self.cpsIN, text='Sócio',\
            background='lightblue', font=('Times New Roman',15,'bold italic'))\
                .place(relx=0.05,rely=0.42)
                
        self.canvas = Canvas(self.cpsIN, width=655, height=10,border=-5)
        self.canvas.place(relx=0.13,rely=0.455)
                
        self.canvas.create_line(-5,0,655,0, fill="darkblue", width=10)
        
        # x,angulo x , y, angulo y

        ###########nome

        Label(self.cpsIN, text='Nome sócio',\
            background='lightblue', font=(10))\
                .place(relx=0.05,rely=0.48)

        self.nomeEntry = Entry(self.cpsIN,\
            textvariable=self.referencias['$nomeContra'],\
                validate='key', validatecommand=(self.cpsIN.register(lambda text: not text.isdecimal()), '%S'))\
                .place(relx=0.05,rely=0.53,relwidth=0.25,relheight=0.05)

        ###########rua

        Label(self.cpsIN, text='Rua',\
            background='lightblue', font=(10))\
                .place(relx=0.35,rely=0.48)

        self.ruaEntry = Entry(self.cpsIN,\
            textvariable=self.referencias['$ruaContra'],\
                validate='key', validatecommand=(self.cpsIN.register(lambda text: not text.isdecimal()), '%S'))\
                .place(relx=0.35,rely=0.53,relwidth=0.20,relheight=0.05)

        ###########Num

        Label(self.cpsIN, text='Num.',\
            background='lightblue', font=(10))\
                .place(relx=0.6,rely=0.48)

        self.numEntry = Entry(self.cpsIN,\
            textvariable=self.referencias['$numContra'],\
                validate='key', validatecommand=(self.cpsIN.register(lambda text: text.isdecimal()), '%S'))\
                .place(relx=0.61,rely=0.53,relwidth=0.05,relheight=0.05)

        ###########bairro

        Label(self.cpsIN, text='Bairro',\
            background='lightblue', font=(10))\
                .place(relx=0.7,rely=0.48)

        self.bairroEntry = Entry(self.cpsIN,\
            textvariable=self.referencias['$bairroContra'],\
                validate='key', validatecommand=(self.cpsIN.register(lambda text: not text.isdecimal()), '%S'))\
                .place(relx=0.7,rely=0.53,relwidth=0.25,relheight=0.05)

        ###########TODO CEP
        
        self.valCEP_Contra = StringVar()

        self.valCEP_Contra.trace_add('write', lambda *args, passed = self.valCEP_Contra:\
            Formater.cep_formater(passed, *args) )

        Label(self.cpsIN, text='CEP',\
            background='lightblue', font=(10))\
                .place(relx=0.05,rely=0.61)
        

        self.CEPEntry = Entry(self.cpsIN, textvariable = self.valCEP_Contra, \
            validate ='key', validatecommand =(self.cpsIN.register(Validator.cep_validator), '%P'))\
                .place(relx=0.05,rely=0.66,relwidth=0.25,relheight=0.05)

        self.referencias['$cepContra'] = self.valCEP_Contra

        ###########TODO RG
        
        self.valRG = StringVar()

        self.valRG.trace_add('write', lambda *args, passed = self.valRG:\
            Formater.rg_formater(passed, *args) )

        Label(self.cpsIN, text='RG',\
            background='lightblue', font=(10))\
                .place(relx=0.35,rely=0.61)
        

        self.CEPEntry = Entry(self.cpsIN, textvariable = self.valRG, \
            validate ='key', validatecommand =(self.cpsIN.register(Validator.rg_validator), '%P'))\
                .place(relx=0.35,rely=0.66,relwidth=0.2,relheight=0.05)

        self.referencias['$rgContra'] = self.valRG

        ###########Org. Emissor

        Label(self.cpsIN, text='Org.',\
            background='lightblue', font=(10))\
                .place(relx=0.9,rely=0.61)

        self.sspEntry = Entry(self.cpsIN,\
            textvariable=self.referencias['$emissorContra'],\
                validate='key', validatecommand=(self.cpsIN.register(lambda text: not text.isdecimal()), '%S'))\
                .place(relx=0.9,rely=0.66,relwidth=0.05,relheight=0.05)

        ###########TODO CPF

        self.valCPF = StringVar()

        self.valCPF.trace_add('write', lambda *args, passed = self.valCPF:\
            Formater.cpf_formater(passed, *args) )

        Label(self.cpsIN, text='CPF',\
            background='lightblue', font=(10))\
                .place(relx=0.6,rely=0.61)
        

        Entry(self.cpsIN, textvariable = self.valCPF, \
            validate ='key', validatecommand =(self.cpsIN.register(Validator.cpf_validator), '%P'))\
                .place(relx=0.61,rely=0.66,relwidth=0.25,relheight=0.05)

        self.referencias['$cpfContra'] = self.valCPF

        ###########Estado Civil

        Label(self.cpsIN, text='Estado Civil',\
            background='lightblue', font=(10))\
                .place(relx=0.35,rely=0.72)

        self.estadoEntry = StringVar(self.cpsIN)

        self.estadoEntryOpt = ('solteiro(a)', 'casado(a)','divorsiado(a)','viuvo(a)')

        self.estadoEntry.set('solteiro(a)')

        self.popup = OptionMenu(self.cpsIN, self.estadoEntry, *self.estadoEntryOpt)\
            .place(relx=0.35,rely=0.77,relwidth=0.2,relheight=0.06)

        self.referencias['$estadoCivilContra'] = self.estadoEntry

        ###########Complemento

        Label(self.cpsIN, text='Complemento',\
            background='lightblue', font=(10))\
                .place(relx=0.6,rely=0.73)

        self.complementoEntry = Entry(self.cpsIN,\
            textvariable=self.referencias['$compleContra'])\
                .place(relx=0.61,rely=0.78,relwidth=0.35,relheight=0.05)

        #Contrato
        Label(self.cpsIN, text='Contrato',\
            background='lightblue', font=('Times New Roman',15,'bold italic'))\
                .place(relx=0.05,rely=0.81)
                
        self.canvas = Canvas(self.cpsIN, width=625, height=10,border=-5)
        self.canvas.place(relx=0.17,rely=0.845)
                
        self.canvas.create_line(-5,0,625,0, fill="darkblue", width=10)
        
        # x,angulo x , y, angulo y

        ###########TODO Valor pagamento

        self.valPag = StringVar(value='R$ ')

        Label(self.cpsIN, text='Val. Pagamento',\
            background='lightblue', font=(10))\
                .place(relx=0.05,rely=0.88)
        
        self.valEntry = Entry(self.cpsIN, textvariable = self.valPag, )\
                .place(relx=0.05,rely=0.93,relwidth=0.1,relheight=0.05)

        self.referencias['$valPag'] = self.valPag

        ###########TODO Data inicio
        
        self.valDT_inic = StringVar()

        self.valDT_inic.trace_add('write', lambda *args, passed = self.valDT_inic:\
            Formater.date_formater(passed, *args) )

        Label(self.cpsIN, text='Data início',\
            background='lightblue', font=(10))\
                .place(relx=0.25,rely=0.88)
        

        Entry(self.cpsIN, textvariable = self.valDT_inic, \
            validate ='key', validatecommand =(self.cpsIN.register(Validator.date_validator), '%P'))\
                .place(relx=0.25,rely=0.93,relwidth=0.1,relheight=0.05)

        self.referencias['$dtInic'] = self.valDT_inic

        ###########TODO Data vencimento
        
        self.valDT_venc = StringVar()

        self.valDT_venc.trace_add('write', lambda *args, passed = self.valDT_venc:\
            Formater.date_formater(passed, *args) )

        Label(self.cpsIN, text='Data vencimento',\
            background='lightblue', font=(10))\
                .place(relx=0.4,rely=0.88)
        

        Entry(self.cpsIN, textvariable = self.valDT_venc, \
            validate ='key', validatecommand =(self.cpsIN.register(Validator.date_validator), '%P'))\
                .place(relx=0.4,rely=0.93,relwidth=0.1,relheight=0.05)

        self.referencias['$dtVenc'] = self.valDT_venc

        #Botão enviar
        Button(self.cpsIN, text='Gerar CPS',\
            command= lambda: self.alterar_doc(frame_ativo=self.cpsIN))\
                .place(relx=0.61,rely=0.865,relwidth=0.35,relheight=0.12)

    def pageLP(self):
        self.menu.destroy()
        self.doc = Document('./code/CPS\'s/CPS LUCRO PRESUMIDO.docx')

        self.cpsLP = Frame(self.window, bd=4, bg='lightblue')
        self.cpsLP.place(relx=0.05,rely=0.05,relwidth=0.9,relheight=0.9)

        #Titulo
        Label(self.cpsLP, text='Gerador de CPS Lucro Presumido', background='lightblue', font=('arial',17,'bold'))\
            .place(relx=0.3,rely=0.045)
            
        #Logo
        self.logo = PhotoImage(file='./code/imgs/deltaprice_logo-slim.png')
        
        self.logo = self.logo.subsample(5,5)
        
        Label(self.cpsLP, image=self.logo, background='lightblue')\
            .place(relx=0.75,rely=0.01,relwidth=0.12,relheight=0.15)

        #Botão voltar
        Button(self.cpsLP, text='Voltar ao menu',\
            command= lambda: (self.cpsLP.destroy, self.pageMenu()))\
                .place(relx=0,rely=0,relwidth=0.25,relheight=0.06)

        self.referencias = {
            '$nomeEmp' : StringVar(), #strValidator
            '$ruaEmp' : StringVar(), #strValidator
            '$numEmp' : StringVar(), #intValidator
            '$bairroEmp' : StringVar(),  #strValidator
            '$cepEmp' : StringVar(),  #intValidator
            '$rgEmp' : StringVar(),  #rgValidator
            '$sspEmp' : StringVar(),  #sspValidator
            '$cnpjEmp' : StringVar(),  #cpfValidator
            '$nomeContra' : StringVar(), #strValidator
            '$ruaContra' : StringVar(), #strValidator
            '$numContra' : StringVar(), #intValidator
            '$bairroContra' : StringVar(),  #strValidator
            '$cepContra' : StringVar(),  #intValidator
            '$rgContra' : StringVar(),  #rgValidator
            '$emissorContra' : StringVar(),  #sspValidator
            '$cpfContra' : StringVar(),  #cpfValidator
            '$estadoCivilContra' : StringVar(), #strValidator
            "$compleEmp" : StringVar(), #strValidator
            "$compleContra" : StringVar(), #strValidator
            "$dtVenc" : StringVar(),  #dateValidator
            "$valPag" : StringVar(), #intValidator
            "$dtInic" : StringVar()  #dateValidator
        }

        #Labels e Entrys
        #Empresa
        Label(self.cpsLP, text='Empresa',\
            background='lightblue', font=('Times New Roman',15,'bold italic'))\
                .place(relx=0.05,rely=0.115)
                
        self.canvas = Canvas(self.cpsLP, width=625, height=10, background='darkblue',border=-5)
        self.canvas.place(relx=0.17,rely=0.15)
                
        self.canvas.create_line(-5,0,625,0, fill="darkblue", width=10)
        
        # x,angulo x , y, angulo y
                
        ###########nome empresa

        Label(self.cpsLP, text='Nome empresa',\
            background='lightblue', font=(10))\
                .place(relx=0.05,rely=0.18)

        self.nomeEmpEntry = Entry(self.cpsLP,\
            textvariable=self.referencias['$nomeEmp'],\
                validate='key', validatecommand=(self.cpsLP.register(lambda text: not text.isdecimal()), '%S'))\
                .place(relx=0.05,rely=0.23,relwidth=0.25,relheight=0.05)

        ###########rua

        Label(self.cpsLP, text='Rua',\
            background='lightblue', font=(10))\
                .place(relx=0.35,rely=0.18)

        self.ruaEmpEntry = Entry(self.cpsLP,\
            textvariable=self.referencias['$ruaEmp'],\
                validate='key', validatecommand=(self.cpsLP.register(lambda text: not text.isdecimal()), '%S'))\
                .place(relx=0.35,rely=0.23,relwidth=0.20,relheight=0.05)

        ###########Num

        Label(self.cpsLP, text='Num.',\
            background='lightblue', font=(10))\
                .place(relx=0.6,rely=0.18)

        self.numEmpEntry = Entry(self.cpsLP,\
            textvariable=self.referencias['$numEmp'],\
                validate='key', validatecommand=(self.cpsLP.register(lambda text: text.isdecimal()), '%S'))\
                    .place(relx=0.61,rely=0.23,relwidth=0.05,relheight=0.05)

        ###########bairro

        Label(self.cpsLP, text='Bairro',\
            background='lightblue', font=(10))\
                .place(relx=0.7,rely=0.18)

        self.bairroEmpEntry = Entry(self.cpsLP,\
            textvariable=self.referencias['$bairroEmp'],\
                validate='key', validatecommand=(self.cpsLP.register(lambda text: not text.isdecimal()), '%S'))\
                .place(relx=0.7,rely=0.23,relwidth=0.25,relheight=0.05)

        ########### CEP Empre

        self.valCEP_Empre = StringVar()

        self.valCEP_Empre.trace_add('write', lambda *args, passed = self.valCEP_Empre:\
            Formater.cep_formater(passed, *args) )

        Label(self.cpsLP, text='CEP',\
            background='lightblue', font=(10))\
                .place(relx=0.05,rely=0.31)
        

        self.CEPEntry = Entry(self.cpsLP, textvariable = self.valCEP_Empre, \
            validate ='key', validatecommand =(self.cpsLP.register(Validator.cep_validator), '%P'))\
                .place(relx=0.05,rely=0.36,relwidth=0.25,relheight=0.05)

        self.referencias['$cepEmp'] = self.valCEP_Empre

        ###########TODO CNPJ
        
        self.valCNPJ = StringVar()

        self.valCNPJ.trace_add('write', lambda *args, passed = self.valCNPJ:\
            Formater.cnpj_formater(passed, *args) )

        Label(self.cpsLP, text='CNPJ',\
            background='lightblue', font=(10))\
                .place(relx=0.35,rely=0.31)
        

        self.CEPEntry = Entry(self.cpsLP, textvariable = self.valCNPJ, \
            validate ='key', validatecommand =(self.cpsLP.register(Validator.cnpj_validator), '%P'))\
                .place(relx=0.35,rely=0.36,relwidth=0.2,relheight=0.05)

        self.referencias['$cnpjEmp'] = self.valCNPJ
                
        ###########Complemento

        Label(self.cpsLP, text='Complemento',\
            background='lightblue', font=(10))\
                .place(relx=0.6,rely=0.31)

        self.complementoEntry = Entry(self.cpsLP,\
            textvariable=self.referencias['$compleEmp'])\
                .place(relx=0.61,rely=0.36,relwidth=0.35,relheight=0.05)
        
        #Socio
        Label(self.cpsLP, text='Sócio',\
            background='lightblue', font=('Times New Roman',15,'bold italic'))\
                .place(relx=0.05,rely=0.42)
                
        self.canvas = Canvas(self.cpsLP, width=655, height=10,border=-5)
        self.canvas.place(relx=0.13,rely=0.455)
                
        self.canvas.create_line(-5,0,655,0, fill="darkblue", width=10)
        
        # x,angulo x , y, angulo y

        ###########nome

        Label(self.cpsLP, text='Nome sócio',\
            background='lightblue', font=(10))\
                .place(relx=0.05,rely=0.48)

        self.nomeEntry = Entry(self.cpsLP,\
            textvariable=self.referencias['$nomeContra'],\
                validate='key', validatecommand=(self.cpsLP.register(lambda text: not text.isdecimal()), '%S'))\
                .place(relx=0.05,rely=0.53,relwidth=0.25,relheight=0.05)

        ###########rua

        Label(self.cpsLP, text='Rua',\
            background='lightblue', font=(10))\
                .place(relx=0.35,rely=0.48)

        self.ruaEntry = Entry(self.cpsLP,\
            textvariable=self.referencias['$ruaContra'],\
                validate='key', validatecommand=(self.cpsLP.register(lambda text: not text.isdecimal()), '%S'))\
                .place(relx=0.35,rely=0.53,relwidth=0.20,relheight=0.05)

        ###########Num

        Label(self.cpsLP, text='Num.',\
            background='lightblue', font=(10))\
                .place(relx=0.6,rely=0.48)

        self.numEntry = Entry(self.cpsLP,\
            textvariable=self.referencias['$numContra'],\
                validate='key', validatecommand=(self.cpsLP.register(lambda text: text.isdecimal()), '%S'))\
                .place(relx=0.61,rely=0.53,relwidth=0.05,relheight=0.05)

        ###########bairro

        Label(self.cpsLP, text='Bairro',\
            background='lightblue', font=(10))\
                .place(relx=0.7,rely=0.48)

        self.bairroEntry = Entry(self.cpsLP,\
            textvariable=self.referencias['$bairroContra'],\
                validate='key', validatecommand=(self.cpsLP.register(lambda text: not text.isdecimal()), '%S'))\
                .place(relx=0.7,rely=0.53,relwidth=0.25,relheight=0.05)

        ###########TODO CEP
        
        self.valCEP_Contra = StringVar()

        self.valCEP_Contra.trace_add('write', lambda *args, passed = self.valCEP_Contra:\
            Formater.cep_formater(passed, *args) )

        Label(self.cpsLP, text='CEP',\
            background='lightblue', font=(10))\
                .place(relx=0.05,rely=0.61)
        

        self.CEPEntry = Entry(self.cpsLP, textvariable = self.valCEP_Contra, \
            validate ='key', validatecommand =(self.cpsLP.register(Validator.cep_validator), '%P'))\
                .place(relx=0.05,rely=0.66,relwidth=0.25,relheight=0.05)

        self.referencias['$cepContra'] = self.valCEP_Contra

        ###########TODO RG
        
        self.valRG = StringVar()

        self.valRG.trace_add('write', lambda *args, passed = self.valRG:\
            Formater.rg_formater(passed, *args) )

        Label(self.cpsLP, text='RG',\
            background='lightblue', font=(10))\
                .place(relx=0.35,rely=0.61)
        

        self.CEPEntry = Entry(self.cpsLP, textvariable = self.valRG, \
            validate ='key', validatecommand =(self.cpsLP.register(Validator.rg_validator), '%P'))\
                .place(relx=0.35,rely=0.66,relwidth=0.2,relheight=0.05)

        self.referencias['$rgContra'] = self.valRG

        ###########Org. Emissor

        Label(self.cpsLP, text='Org.',\
            background='lightblue', font=(10))\
                .place(relx=0.9,rely=0.61)

        self.sspEntry = Entry(self.cpsLP,\
            textvariable=self.referencias['$emissorContra'],\
                validate='key', validatecommand=(self.cpsLP.register(lambda text: not text.isdecimal()), '%S'))\
                .place(relx=0.9,rely=0.66,relwidth=0.05,relheight=0.05)

        ###########TODO CPF

        self.valCPF = StringVar()

        self.valCPF.trace_add('write', lambda *args, passed = self.valCPF:\
            Formater.cpf_formater(passed, *args) )

        Label(self.cpsLP, text='CPF',\
            background='lightblue', font=(10))\
                .place(relx=0.6,rely=0.61)
        

        Entry(self.cpsLP, textvariable = self.valCPF, \
            validate ='key', validatecommand =(self.cpsLP.register(Validator.cpf_validator), '%P'))\
                .place(relx=0.61,rely=0.66,relwidth=0.25,relheight=0.05)

        self.referencias['$cpfContra'] = self.valCPF

        ###########Estado Civil

        Label(self.cpsLP, text='Estado Civil',\
            background='lightblue', font=(10))\
                .place(relx=0.35,rely=0.72)

        self.estadoEntry = StringVar(self.cpsLP)

        self.estadoEntryOpt = ('solteiro(a)', 'casado(a)','divorsiado(a)','viuvo(a)')

        self.estadoEntry.set('solteiro(a)')

        self.popup = OptionMenu(self.cpsLP, self.estadoEntry, *self.estadoEntryOpt)\
            .place(relx=0.35,rely=0.77,relwidth=0.2,relheight=0.06)

        self.referencias['$estadoCivilContra'] = self.estadoEntry

        ###########Complemento

        Label(self.cpsLP, text='Complemento',\
            background='lightblue', font=(10))\
                .place(relx=0.6,rely=0.73)

        self.complementoEntry = Entry(self.cpsLP,\
            textvariable=self.referencias['$compleContra'])\
                .place(relx=0.61,rely=0.78,relwidth=0.35,relheight=0.05)

        #Contrato
        Label(self.cpsLP, text='Contrato',\
            background='lightblue', font=('Times New Roman',15,'bold italic'))\
                .place(relx=0.05,rely=0.81)
                
        self.canvas = Canvas(self.cpsLP, width=625, height=10,border=-5)
        self.canvas.place(relx=0.17,rely=0.845)
                
        self.canvas.create_line(-5,0,625,0, fill="darkblue", width=10)
        
        # x,angulo x , y, angulo y

        ###########TODO Valor pagamento

        self.valPag = StringVar(value='R$ ')

        Label(self.cpsLP, text='Val. Pagamento',\
            background='lightblue', font=(10))\
                .place(relx=0.05,rely=0.88)
        
        self.valEntry = Entry(self.cpsLP, textvariable = self.valPag, )\
                .place(relx=0.05,rely=0.93,relwidth=0.1,relheight=0.05)

        self.referencias['$valPag'] = self.valPag

        ###########TODO Data inicio
        
        self.valDT_inic = StringVar()

        self.valDT_inic.trace_add('write', lambda *args, passed = self.valDT_inic:\
            Formater.date_formater(passed, *args) )

        Label(self.cpsLP, text='Data início',\
            background='lightblue', font=(10))\
                .place(relx=0.25,rely=0.88)
        

        Entry(self.cpsLP, textvariable = self.valDT_inic, \
            validate ='key', validatecommand =(self.cpsLP.register(Validator.date_validator), '%P'))\
                .place(relx=0.25,rely=0.93,relwidth=0.1,relheight=0.05)

        self.referencias['$dtInic'] = self.valDT_inic

        ###########TODO Data vencimento
        
        self.valDT_venc = StringVar()

        self.valDT_venc.trace_add('write', lambda *args, passed = self.valDT_venc:\
            Formater.date_formater(passed, *args) )

        Label(self.cpsLP, text='Data vencimento',\
            background='lightblue', font=(10))\
                .place(relx=0.4,rely=0.88)
        

        Entry(self.cpsLP, textvariable = self.valDT_venc, \
            validate ='key', validatecommand =(self.cpsLP.register(Validator.date_validator), '%P'))\
                .place(relx=0.4,rely=0.93,relwidth=0.1,relheight=0.05)

        self.referencias['$dtVenc'] = self.valDT_venc

        #Botão enviar
        Button(self.cpsLP, text='Gerar CPS',\
            command= lambda: self.alterar_doc(frame_ativo=self.cpsLP))\
                .place(relx=0.61,rely=0.865,relwidth=0.35,relheight=0.12)

    def pageSN(self):
        self.menu.destroy()
        self.doc = Document('./code/CPS\'s/CPS SIMPLES NACIONAL.docx')

        self.cpsSN = Frame(self.window, bd=4, bg='lightblue')
        self.cpsSN.place(relx=0.05,rely=0.05,relwidth=0.9,relheight=0.9)

        #Titulo
        Label(self.cpsSN, text='Gerador de CPS Simples Nacional', background='lightblue', font=('arial',17,'bold'))\
            .place(relx=0.3,rely=0.045)
            
        #Logo
        self.logo = PhotoImage(file='./code/imgs/deltaprice_logo-slim.png')
        
        self.logo = self.logo.subsample(5,5)
        
        Label(self.cpsSN, image=self.logo, background='lightblue')\
            .place(relx=0.75,rely=0.01,relwidth=0.12,relheight=0.15)

        #Botão voltar
        Button(self.cpsSN, text='Voltar ao menu',\
            command= lambda: (self.cpsSN.destroy, self.pageMenu()))\
                .place(relx=0,rely=0,relwidth=0.25,relheight=0.06)

        self.referencias = {
            '$nomeEmp' : StringVar(), #strValidator
            '$ruaEmp' : StringVar(), #strValidator
            '$numEmp' : StringVar(), #intValidator
            '$bairroEmp' : StringVar(),  #strValidator
            '$cepEmp' : StringVar(),  #intValidator
            '$rgEmp' : StringVar(),  #rgValidator
            '$sspEmp' : StringVar(),  #sspValidator
            '$cnpjEmp' : StringVar(),  #cpfValidator
            '$nomeContra' : StringVar(), #strValidator
            '$ruaContra' : StringVar(), #strValidator
            '$numContra' : StringVar(), #intValidator
            '$bairroContra' : StringVar(),  #strValidator
            '$cepContra' : StringVar(),  #intValidator
            '$rgContra' : StringVar(),  #rgValidator
            '$emissorContra' : StringVar(),  #sspValidator
            '$cpfContra' : StringVar(),  #cpfValidator
            '$estadoCivilContra' : StringVar(), #strValidator
            "$compleEmp" : StringVar(), #strValidator
            "$compleContra" : StringVar(), #strValidator
            "$dtVenc" : StringVar(),  #dateValidator
            "$valPag" : StringVar(), #intValidator
            "$dtInic" : StringVar()  #dateValidator
        }

        #Labels e Entrys
        #Empresa
        Label(self.cpsSN, text='Empresa',\
            background='lightblue', font=('Times New Roman',15,'bold italic'))\
                .place(relx=0.05,rely=0.115)
                
        self.canvas = Canvas(self.cpsSN, width=625, height=10, background='darkblue',border=-5)
        self.canvas.place(relx=0.17,rely=0.15)
                
        self.canvas.create_line(-5,0,625,0, fill="darkblue", width=10)
        
        # x,angulo x , y, angulo y
                
        ###########nome empresa

        Label(self.cpsSN, text='Nome empresa',\
            background='lightblue', font=(10))\
                .place(relx=0.05,rely=0.18)

        self.nomeEmpEntry = Entry(self.cpsSN,\
            textvariable=self.referencias['$nomeEmp'],\
                validate='key', validatecommand=(self.cpsSN.register(lambda text: not text.isdecimal()), '%S'))\
                .place(relx=0.05,rely=0.23,relwidth=0.25,relheight=0.05)

        ###########rua

        Label(self.cpsSN, text='Rua',\
            background='lightblue', font=(10))\
                .place(relx=0.35,rely=0.18)

        self.ruaEmpEntry = Entry(self.cpsSN,\
            textvariable=self.referencias['$ruaEmp'],\
                validate='key', validatecommand=(self.cpsSN.register(lambda text: not text.isdecimal()), '%S'))\
                .place(relx=0.35,rely=0.23,relwidth=0.20,relheight=0.05)

        ###########Num

        Label(self.cpsSN, text='Num.',\
            background='lightblue', font=(10))\
                .place(relx=0.6,rely=0.18)

        self.numEmpEntry = Entry(self.cpsSN,\
            textvariable=self.referencias['$numEmp'],\
                validate='key', validatecommand=(self.cpsSN.register(lambda text: text.isdecimal()), '%S'))\
                    .place(relx=0.61,rely=0.23,relwidth=0.05,relheight=0.05)

        ###########bairro

        Label(self.cpsSN, text='Bairro',\
            background='lightblue', font=(10))\
                .place(relx=0.7,rely=0.18)

        self.bairroEmpEntry = Entry(self.cpsSN,\
            textvariable=self.referencias['$bairroEmp'],\
                validate='key', validatecommand=(self.cpsSN.register(lambda text: not text.isdecimal()), '%S'))\
                .place(relx=0.7,rely=0.23,relwidth=0.25,relheight=0.05)

        ########### CEP Empre

        self.valCEP_Empre = StringVar()

        self.valCEP_Empre.trace_add('write', lambda *args, passed = self.valCEP_Empre:\
            Formater.cep_formater(passed, *args) )

        Label(self.cpsSN, text='CEP',\
            background='lightblue', font=(10))\
                .place(relx=0.05,rely=0.31)
        

        self.CEPEntry = Entry(self.cpsSN, textvariable = self.valCEP_Empre, \
            validate ='key', validatecommand =(self.cpsSN.register(Validator.cep_validator), '%P'))\
                .place(relx=0.05,rely=0.36,relwidth=0.25,relheight=0.05)

        self.referencias['$cepEmp'] = self.valCEP_Empre

        ###########TODO CNPJ
        
        self.valCNPJ = StringVar()

        self.valCNPJ.trace_add('write', lambda *args, passed = self.valCNPJ:\
            Formater.cnpj_formater(passed, *args) )

        Label(self.cpsSN, text='CNPJ',\
            background='lightblue', font=(10))\
                .place(relx=0.35,rely=0.31)
        

        self.CEPEntry = Entry(self.cpsSN, textvariable = self.valCNPJ, \
            validate ='key', validatecommand =(self.cpsSN.register(Validator.cnpj_validator), '%P'))\
                .place(relx=0.35,rely=0.36,relwidth=0.2,relheight=0.05)

        self.referencias['$cnpjEmp'] = self.valCNPJ
                
        ###########Complemento

        Label(self.cpsSN, text='Complemento',\
            background='lightblue', font=(10))\
                .place(relx=0.6,rely=0.31)

        self.complementoEntry = Entry(self.cpsSN,\
            textvariable=self.referencias['$compleEmp'])\
                .place(relx=0.61,rely=0.36,relwidth=0.35,relheight=0.05)
        
        #Socio
        Label(self.cpsSN, text='Sócio',\
            background='lightblue', font=('Times New Roman',15,'bold italic'))\
                .place(relx=0.05,rely=0.42)
                
        self.canvas = Canvas(self.cpsSN, width=655, height=10,border=-5)
        self.canvas.place(relx=0.13,rely=0.455)
                
        self.canvas.create_line(-5,0,655,0, fill="darkblue", width=10)
        
        # x,angulo x , y, angulo y

        ###########nome

        Label(self.cpsSN, text='Nome sócio',\
            background='lightblue', font=(10))\
                .place(relx=0.05,rely=0.48)

        self.nomeEntry = Entry(self.cpsSN,\
            textvariable=self.referencias['$nomeContra'],\
                validate='key', validatecommand=(self.cpsSN.register(lambda text: not text.isdecimal()), '%S'))\
                .place(relx=0.05,rely=0.53,relwidth=0.25,relheight=0.05)

        ###########rua

        Label(self.cpsSN, text='Rua',\
            background='lightblue', font=(10))\
                .place(relx=0.35,rely=0.48)

        self.ruaEntry = Entry(self.cpsSN,\
            textvariable=self.referencias['$ruaContra'],\
                validate='key', validatecommand=(self.cpsSN.register(lambda text: not text.isdecimal()), '%S'))\
                .place(relx=0.35,rely=0.53,relwidth=0.20,relheight=0.05)

        ###########Num

        Label(self.cpsSN, text='Num.',\
            background='lightblue', font=(10))\
                .place(relx=0.6,rely=0.48)

        self.numEntry = Entry(self.cpsSN,\
            textvariable=self.referencias['$numContra'],\
                validate='key', validatecommand=(self.cpsSN.register(lambda text: text.isdecimal()), '%S'))\
                .place(relx=0.61,rely=0.53,relwidth=0.05,relheight=0.05)

        ###########bairro

        Label(self.cpsSN, text='Bairro',\
            background='lightblue', font=(10))\
                .place(relx=0.7,rely=0.48)

        self.bairroEntry = Entry(self.cpsSN,\
            textvariable=self.referencias['$bairroContra'],\
                validate='key', validatecommand=(self.cpsSN.register(lambda text: not text.isdecimal()), '%S'))\
                .place(relx=0.7,rely=0.53,relwidth=0.25,relheight=0.05)

        ###########TODO CEP
        
        self.valCEP_Contra = StringVar()

        self.valCEP_Contra.trace_add('write', lambda *args, passed = self.valCEP_Contra:\
            Formater.cep_formater(passed, *args) )

        Label(self.cpsSN, text='CEP',\
            background='lightblue', font=(10))\
                .place(relx=0.05,rely=0.61)
        

        self.CEPEntry = Entry(self.cpsSN, textvariable = self.valCEP_Contra, \
            validate ='key', validatecommand =(self.cpsSN.register(Validator.cep_validator), '%P'))\
                .place(relx=0.05,rely=0.66,relwidth=0.25,relheight=0.05)

        self.referencias['$cepContra'] = self.valCEP_Contra

        ###########TODO RG
        
        self.valRG = StringVar()

        self.valRG.trace_add('write', lambda *args, passed = self.valRG:\
            Formater.rg_formater(passed, *args) )

        Label(self.cpsSN, text='RG',\
            background='lightblue', font=(10))\
                .place(relx=0.35,rely=0.61)
        

        self.CEPEntry = Entry(self.cpsSN, textvariable = self.valRG, \
            validate ='key', validatecommand =(self.cpsSN.register(Validator.rg_validator), '%P'))\
                .place(relx=0.35,rely=0.66,relwidth=0.2,relheight=0.05)

        self.referencias['$rgContra'] = self.valRG

        ###########Org. Emissor

        Label(self.cpsSN, text='Org.',\
            background='lightblue', font=(10))\
                .place(relx=0.9,rely=0.61)

        self.sspEntry = Entry(self.cpsSN,\
            textvariable=self.referencias['$emissorContra'],\
                validate='key', validatecommand=(self.cpsSN.register(lambda text: not text.isdecimal()), '%S'))\
                .place(relx=0.9,rely=0.66,relwidth=0.05,relheight=0.05)

        ###########TODO CPF

        self.valCPF = StringVar()

        self.valCPF.trace_add('write', lambda *args, passed = self.valCPF:\
            Formater.cpf_formater(passed, *args) )

        Label(self.cpsSN, text='CPF',\
            background='lightblue', font=(10))\
                .place(relx=0.6,rely=0.61)
        

        Entry(self.cpsSN, textvariable = self.valCPF, \
            validate ='key', validatecommand =(self.cpsSN.register(Validator.cpf_validator), '%P'))\
                .place(relx=0.61,rely=0.66,relwidth=0.25,relheight=0.05)

        self.referencias['$cpfContra'] = self.valCPF

        ###########Estado Civil

        Label(self.cpsSN, text='Estado Civil',\
            background='lightblue', font=(10))\
                .place(relx=0.35,rely=0.72)

        self.estadoEntry = StringVar(self.cpsSN)

        self.estadoEntryOpt = ('solteiro(a)', 'casado(a)','divorsiado(a)','viuvo(a)')

        self.estadoEntry.set('solteiro(a)')

        self.popup = OptionMenu(self.cpsSN, self.estadoEntry, *self.estadoEntryOpt)\
            .place(relx=0.35,rely=0.77,relwidth=0.2,relheight=0.06)

        self.referencias['$estadoCivilContra'] = self.estadoEntry

        ###########Complemento

        Label(self.cpsSN, text='Complemento',\
            background='lightblue', font=(10))\
                .place(relx=0.6,rely=0.73)

        self.complementoEntry = Entry(self.cpsSN,\
            textvariable=self.referencias['$compleContra'])\
                .place(relx=0.61,rely=0.78,relwidth=0.35,relheight=0.05)

        #Contrato
        Label(self.cpsSN, text='Contrato',\
            background='lightblue', font=('Times New Roman',15,'bold italic'))\
                .place(relx=0.05,rely=0.81)
                
        self.canvas = Canvas(self.cpsSN, width=625, height=10,border=-5)
        self.canvas.place(relx=0.17,rely=0.845)
                
        self.canvas.create_line(-5,0,625,0, fill="darkblue", width=10)
        
        # x,angulo x , y, angulo y

        ###########TODO Valor pagamento

        self.valPag = StringVar(value='R$ ')

        Label(self.cpsSN, text='Val. Pagamento',\
            background='lightblue', font=(10))\
                .place(relx=0.05,rely=0.88)
        
        self.valEntry = Entry(self.cpsSN, textvariable = self.valPag, )\
                .place(relx=0.05,rely=0.93,relwidth=0.1,relheight=0.05)

        self.referencias['$valPag'] = self.valPag

        ###########TODO Data inicio
        
        self.valDT_inic = StringVar()

        self.valDT_inic.trace_add('write', lambda *args, passed = self.valDT_inic:\
            Formater.date_formater(passed, *args) )

        Label(self.cpsSN, text='Data início',\
            background='lightblue', font=(10))\
                .place(relx=0.25,rely=0.88)
        

        Entry(self.cpsSN, textvariable = self.valDT_inic, \
            validate ='key', validatecommand =(self.cpsSN.register(Validator.date_validator), '%P'))\
                .place(relx=0.25,rely=0.93,relwidth=0.1,relheight=0.05)

        self.referencias['$dtInic'] = self.valDT_inic

        ###########TODO Data vencimento
        
        self.valDT_venc = StringVar()

        self.valDT_venc.trace_add('write', lambda *args, passed = self.valDT_venc:\
            Formater.date_formater(passed, *args) )

        Label(self.cpsSN, text='Data vencimento',\
            background='lightblue', font=(10))\
                .place(relx=0.4,rely=0.88)
        

        Entry(self.cpsSN, textvariable = self.valDT_venc, \
            validate ='key', validatecommand =(self.cpsSN.register(Validator.date_validator), '%P'))\
                .place(relx=0.4,rely=0.93,relwidth=0.1,relheight=0.05)

        self.referencias['$dtVenc'] = self.valDT_venc

        #Botão enviar
        Button(self.cpsSN, text='Gerar CPS',\
            command= lambda: self.alterar_doc(frame_ativo=self.cpsSN))\
                .place(relx=0.61,rely=0.865,relwidth=0.35,relheight=0.12)

App()